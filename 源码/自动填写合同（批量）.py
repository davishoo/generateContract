from openpyxl import load_workbook
from docx import Document
from os import listdir

def replace_text(old_text, new_text):
    all_paragraphs = document.paragraphs #读取所有的自然段
    for paragraph in all_paragraphs:
        for run in paragraph.runs: #循环读取所有的run，并进行新旧文本的替换
            run_text = run.text.replace(old_text, new_text)
            run.text = run_text    
    all_tables = document.tables #读取所有的表格
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells: #循环读取表格中所有的cells，并进行新旧文本的替换
                cell_text = cell.text.replace(old_text, new_text)
                cell.text = cell_text

for file in listdir():
    if '.xlsx' in file:
        xlsx_name = file
    if '.docx' in file:
        docx_name = file

wb = load_workbook(xlsx_name)
sheetx0 = wb.sheetnames
sheetx = wb[sheetx0[0]]



for col in range(2,sheetx.max_column+1): #合同要素excel中逐列循环
    document = Document(docx_name)
    if sheetx.cell(row=1,column=col).value!=None: #openpyxl在使用sheetx.max_column时可能会读取到空的单元格，这里进行剔除
        for l in range(1,sheetx.max_row+1): #合同要素excel中逐行循环
            old_text = sheetx.cell(row=l,column=1).value #合同要素excel中对第一列逐行读取编号
            new_text = sheetx.cell(row=l,column=col).value #合同要素excel中对循环的当前列逐行读取新要素
            replace_text(str(old_text),str(new_text)) #进行替换
            filename = str(sheetx.cell(row=1,column=col).value) #定义文件名为当前列第一行的内容
        document.save("%s.docx"%(filename)) #按定义的文件名进行保存
print('ok')
